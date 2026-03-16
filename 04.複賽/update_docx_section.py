import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

INPUT = r'C:\Users\226178\Desktop\GitHub\startup-asvda\04.複賽\B0274商業計劃書v1.docx'
OUTPUT = r'C:\Users\226178\Desktop\GitHub\startup-asvda\04.複賽\B0274商業計劃書v1_updated.docx'

doc = Document(INPUT)
body = doc.element.body
elements = list(body)

# Find element indices to remove (103-112: from "1. 單位經濟模型" to 毛利率 content)
START_IDX = 103  # "1. 單位經濟模型"
END_IDX = 112    # Last element before "3. 通路與品牌策略"

# Store reference to the element BEFORE which we insert (element 113 = "3. 通路與品牌策略")
insert_before = elements[113]

# Remove old elements
for idx in range(START_IDX, END_IDX + 1):
    body.remove(elements[idx])

# === Helper functions ===
HEADER_COLOR = '1A5276'
BODY_FONT = '微軟正黑體'
BORDER_COLOR = 'BDC3C7'
TEXT_COLOR = '1B2631'
ACCENT_COLOR = '1A5276'

def make_paragraph(text, bold=False, color=TEXT_COLOR, font_east=BODY_FONT, spacing_before=60, spacing_after=60, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="{spacing_before}" w:after="{spacing_after}" w:line="360" w:lineRule="auto"/><w:jc w:val="both"/></w:pPr></w:p>')
    return p

def add_run(p, text, bold=False, color=TEXT_COLOR, font_east=BODY_FONT, sz=None):
    r = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:eastAsia="{font_east}"/></w:rPr><w:t xml:space="preserve">{escape(text)}</w:t></w:r>')
    rpr = r.find(qn('w:rPr'))
    if bold:
        rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        rpr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
    color_el = parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>')
    rpr.append(color_el)
    if sz:
        rpr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{sz}"/>'))
    p.append(r)
    return r

def escape(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def make_heading3(text):
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="3"/></w:pPr></w:p>')
    # Split "1. xxx" into number and rest
    if '. ' in text and text[0].isdigit():
        num, rest = text.split('. ', 1)
        r1 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:hint="eastAsia"/></w:rPr><w:t>{escape(num)}</w:t></w:r>')
        r2 = parse_xml(f'<w:r {nsdecls("w")}><w:t>{escape(". " + rest)}</w:t></w:r>')
        p.append(r1)
        p.append(r2)
    else:
        r = parse_xml(f'<w:r {nsdecls("w")}><w:t>{escape(text)}</w:t></w:r>')
        p.append(r)
    return p

def make_subheading(text):
    p = make_paragraph(text, bold=True)
    add_run(p, text, bold=True, color=ACCENT_COLOR)
    return p

def make_body_text(parts):
    """parts is list of (text, bold, color) tuples"""
    p = make_paragraph('')
    for text, bold, color in parts:
        add_run(p, text, bold=bold, color=color)
    return p

def make_simple_para(text):
    p = make_paragraph('')
    add_run(p, text)
    return p

def make_bold_para(text):
    p = make_paragraph('')
    add_run(p, text, bold=True)
    return p

def make_table(headers, rows, col_widths=None):
    """Create a formatted table with dark header row"""
    ncols = len(headers)
    total_w = 9026
    if col_widths is None:
        col_widths = [total_w // ncols] * ncols
        col_widths[-1] = total_w - sum(col_widths[:-1])

    tbl_xml = f'<w:tbl {nsdecls("w")}>'
    tbl_xml += f'<w:tblPr><w:tblW w:w="{total_w}" w:type="dxa"/>'
    tbl_xml += '<w:tblBorders>'
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tbl_xml += f'<w:{side} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    tbl_xml += '</w:tblBorders>'
    tbl_xml += '<w:tblCellMar><w:left w:w="10" w:type="dxa"/><w:right w:w="10" w:type="dxa"/></w:tblCellMar>'
    tbl_xml += '<w:tblLook w:val="0000" w:firstRow="0" w:lastRow="0" w:firstColumn="0" w:lastColumn="0" w:noHBand="0" w:noVBand="0"/>'
    tbl_xml += '</w:tblPr><w:tblGrid>'
    for w in col_widths:
        tbl_xml += f'<w:gridCol w:w="{w}"/>'
    tbl_xml += '</w:tblGrid>'

    def make_cell(text, width, is_header=False, bold_text=False):
        fill = HEADER_COLOR if is_header else 'FFFFFF'
        txt_color = 'FFFFFF' if is_header else TEXT_COLOR
        cell_xml = f'<w:tc><w:tcPr><w:tcW w:w="{width}" w:type="dxa"/>'
        cell_xml += '<w:tcBorders>'
        for s in ['top', 'left', 'bottom', 'right']:
            cell_xml += f'<w:{s} w:val="single" w:sz="1" w:space="0" w:color="{BORDER_COLOR}"/>'
        cell_xml += '</w:tcBorders>'
        cell_xml += f'<w:shd w:val="clear" w:color="auto" w:fill="{fill}"/>'
        cell_xml += '<w:tcMar><w:top w:w="60" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:bottom w:w="60" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>'
        cell_xml += '<w:vAlign w:val="center"/></w:tcPr>'
        b_tags = '<w:b/><w:bCs/>' if (is_header or bold_text) else ''
        cell_xml += f'<w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="{BODY_FONT}"/>{b_tags}<w:color w:val="{txt_color}"/><w:szCs w:val="20"/></w:rPr><w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>'
        cell_xml += '</w:tc>'
        return cell_xml

    def make_row(cells_data, is_header=False):
        row_xml = '<w:tr><w:tblPrEx><w:tblCellMar><w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tblCellMar></w:tblPrEx>'
        for i, text in enumerate(cells_data):
            bold = isinstance(text, str) and text.startswith('**') and text.endswith('**')
            if bold:
                text = text[2:-2]
            row_xml += make_cell(text, col_widths[i], is_header=is_header, bold_text=bold)
        row_xml += '</w:tr>'
        return row_xml

    tbl_xml += make_row(headers, is_header=True)
    for row in rows:
        tbl_xml += make_row(row)
    tbl_xml += '</w:tbl>'
    return parse_xml(tbl_xml)

def make_bullet_point(text_parts):
    """text_parts: list of (text, bold) tuples"""
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="40" w:after="40" w:line="360" w:lineRule="auto"/><w:ind w:left="360" w:hanging="360"/><w:jc w:val="both"/></w:pPr></w:p>')
    # Add bullet character
    add_run(p, '• ', bold=False, color=TEXT_COLOR)
    for text, bold in text_parts:
        add_run(p, text, bold=bold, color=TEXT_COLOR)
    return p

def make_blockquote(text):
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="60" w:after="60" w:line="360" w:lineRule="auto"/><w:ind w:left="360"/><w:jc w:val="both"/></w:pPr></w:p>')
    add_run(p, text, bold=False, color='808080')
    return p

# === Build new content ===
new_elements = []

# --- 1. 單位經濟模型 ---
new_elements.append(make_heading3('1. 單位經濟模型'))
new_elements.append(make_subheading('收費機制：Freemium SaaS 訂閱制'))

new_elements.append(make_table(
    ['方案', '月費', '功能'],
    [
        ['免費方案', 'NT$0', '每日 3 次 AI 對話 + 基礎題庫'],
        ['**付費方案**', '**NT$600/月**', '無限 AI 對話 + 完整題庫 + 學習筆記 + 知識圖譜 + 模擬考 + 學習報告'],
        ['**年繳方案**', '**NT$500/月（83 折）**', '同付費方案，年繳一次付清'],
    ],
    [1559, 2941, 4526]
))

new_elements.append(make_body_text([
    ('為什麼 Freemium？', True, ACCENT_COLOR),
    (' 因為學生間的口碑傳播是 EdTech 最強的增長引擎。Duolingo（$7.48 億營收）、RevisionDojo（YC 認可）都驗證了這條路——', False, TEXT_COLOR),
    ('免費用戶不是成本，是行銷資產', True, TEXT_COLOR),
    ('。', False, TEXT_COLOR),
]))

new_elements.append(make_body_text([
    ('付費用戶月均收入（ARPPU）= NT$600/月', True, ACCENT_COLOR),
    ('——單一定價方案，價格透明、決策門檻低。', False, TEXT_COLOR),
]))

new_elements.append(make_simple_para('以 Year 1 目標（15,000 免費用戶、500 付費用戶）計算免費→付費轉換率約 3.3%：'))

new_elements.append(make_bullet_point([
    ('ARPU（全用戶）= 總月營收 ÷ 全部用戶 = NT$600 × 500 ÷ 15,500 ≈ NT$19/月', True),
]))

new_elements.append(make_simple_para('ARPU 看似極低，正是 Freemium 模式的典型特徵——絕大多數用戶免費使用並創造口碑價值，少數付費用戶產生集中營收。隨轉換率從 3.3% 提升至 6%（Year 2 目標），ARPU 將同步倍增。後續 LTV 與 CAC 分析以付費用戶為基準，使用 ARPPU = NT$600。'))

# COGS
new_elements.append(make_subheading('每用戶服務成本（COGS）'))
new_elements.append(make_simple_para('只計入隨用戶數量線性增長的變動成本，人事、行銷等固定支出歸入營運費用（OpEx）：'))

new_elements.append(make_table(
    ['成本項目', '月均/用戶', '說明'],
    [
        ['AI 推理成本', 'NT$80–120', '平均 150 次查詢/月 × NT$0.6/次（成本優化後模型）'],
        ['雲端基礎設施', 'NT$15–25', '伺服器、CDN、資料庫依用戶數分攤'],
        ['**合計 COGS**', '**NT$100–140**', '**取中位數 NT$120**'],
    ],
    [2000, 2200, 4826]
))

new_elements.append(make_bullet_point([('單位貢獻利潤', True), (' = NT$600 − NT$120 = ', False), ('NT$480/月', True)]))
new_elements.append(make_bullet_point([('貢獻利潤率', True), (' = 480 ÷ 600 = ', False), ('80%', True)]))
new_elements.append(make_simple_para('排除人事、行銷等固定支出後，每服務一個付費學生，80% 的收入轉化為貢獻利潤——這是 SaaS 模式的核心優勢。'))

# Customer lifecycle
new_elements.append(make_subheading('客戶生命週期（教育場景修正）'))
new_elements.append(make_simple_para('傳統 SaaS 用 1/月流失率 計算生命週期，但教育市場有天然畢業天花板——國中生最長 3 年、高中生最長 3 年。我們必須用加權留存分析取代無限期假設：'))

new_elements.append(make_table(
    ['學生類型', '占比', '平均留存', '說明'],
    [
        ['短期使用（1 學年）', '60%', '10 個月', '段考/會考衝刺後離開'],
        ['中期使用（2 學年）', '30%', '20 個月', '跨學年持續訂閱'],
        ['長期使用（3 年+）', '10%', '30 個月', '國中→高中續用'],
    ],
    [2200, 1200, 1600, 4026]
))

new_elements.append(make_bullet_point([('加權平均生命週期', True), (' = 0.6×10 + 0.3×20 + 0.1×30 = 6 + 6 + 3 = ', False), ('15 個月', True)]))
new_elements.append(make_bullet_point([('考量寒暑假流失（學期間流失率 15–20%）及升學銜接斷層，取保守值 ', False), ('14 個月', True)]))

new_elements.append(make_blockquote('流失率邏輯：學期內月流失率約 3–4%（考試動機強），但寒暑假流失率跳升至 15–20%。年均月流失率約 7%，但因畢業天花板，不能簡單用 1/7% = 14.3 個月做無限推導。'))

# LTV
new_elements.append(make_subheading('LTV 計算'))
new_elements.append(make_bullet_point([('LTV = 單位貢獻利潤 × 平均生命週期 = NT$480 × 14 = NT$6,720', True)]))

# CAC
new_elements.append(make_subheading('CAC 分渠道分析'))

new_elements.append(make_table(
    ['獲客渠道', '預估 CAC', '占比', '說明'],
    [
        ['免費轉付費（PLG）', 'NT$0–200', '40%', '免費用戶自然轉換'],
        ['口碑推薦', 'NT$300–500', '25%', '推薦獎勵機制（贈送付費天數）'],
        ['KOL / 社群行銷', 'NT$800–1,200', '20%', '數學補教 YouTuber、IG 學習帳號合作'],
        ['校園推廣', 'NT$500–800', '15%', '教育展、學校免費試用轉化'],
    ],
    [2200, 1800, 1200, 3826]
))

new_elements.append(make_bullet_point([('加權平均 CAC', True), (' = 0.4×100 + 0.25×400 + 0.2×1,000 + 0.15×650 ≈ ', False), ('NT$440', True)]))
new_elements.append(make_bullet_point([('取 ', False), ('CAC ≈ NT$500', True), ('（含渠道測試與優化成本）', False)]))

# Health metrics
new_elements.append(make_subheading('單位經濟健康度'))

new_elements.append(make_table(
    ['指標', '數值', '業界基準', '判斷'],
    [
        ['LTV', 'NT$6,720', '—', '—'],
        ['CAC', 'NT$500', '—', '—'],
        ['**LTV/CAC**', '**13.4:1**', '>3:1', '✅ 極健康'],
        ['**CAC 回收期**', '**1.0 個月**', '<12 個月', '✅ 極快'],
        ['貢獻利潤率', '80%', '>70%', '✅ SaaS 健康'],
    ],
    [2200, 2200, 2200, 2426]
))

new_elements.append(make_simple_para('LTV/CAC 13.4:1 看似極高，原因在於 PLG + 口碑佔獲客 65%——EdTech 的學生社交傳播天然壓低 CAC。這也意味著：若未來需要加大付費渠道投放，CAC 會上升，但仍有充足的利潤緩衝空間。'))

# Sensitivity
new_elements.append(make_subheading('敏感度分析'))

new_elements.append(make_table(
    ['情境', 'ARPPU', '生命週期', 'LTV', 'CAC', 'LTV/CAC'],
    [
        ['🟢 樂觀', 'NT$600', '18 個月', 'NT$8,640', 'NT$400', '21.6:1'],
        ['🟡 基準', 'NT$600', '14 個月', 'NT$6,720', 'NT$500', '13.4:1'],
        ['🔴 悲觀', 'NT$600', '10 個月', 'NT$4,800', 'NT$700', '6.9:1'],
    ],
    [1400, 1400, 1400, 1600, 1600, 1626]
))

new_elements.append(make_simple_para('ARPPU 固定為 NT$600（單一定價方案），敏感度主要來自留存時間與獲客成本。即使在悲觀情境下，LTV/CAC 仍達 6.9:1，遠高於 3:1 健康線。單位經濟模型具備足夠的安全邊際。'))

# --- 2. 毛利率 ---
new_elements.append(make_heading3('2. 毛利率'))
new_elements.append(make_simple_para('Mathify Labs 採用 SaaS 模式，成本結構可清楚區分為變動成本（COGS）與固定營運費用（OpEx）：'))

new_elements.append(make_subheading('變動成本（隨用戶數線性增長）'))

new_elements.append(make_table(
    ['成本項目', '月均/付費用戶', '占營收比', '說明'],
    [
        ['AI 推理成本', 'NT$80–120', '13–20%', 'LLM API 調用，使用成本優化模型'],
        ['雲端基礎設施', 'NT$15–25', '3–4%', '伺服器、CDN、資料庫依用戶數分攤'],
        ['**COGS 合計**', '**NT$100–140**', '**17–23%**', ''],
    ],
    [2200, 2200, 1800, 2826]
))

new_elements.append(make_bold_para('毛利率 = (ARPPU − COGS) ÷ ARPPU = (600 − 120) ÷ 600 ≈ 80%'))

new_elements.append(make_subheading('固定營運費用（不隨用戶數變動）'))

new_elements.append(make_table(
    ['費用項目', 'Year 1 月均', '說明'],
    [
        ['人事成本', '~NT$6 萬', '首位工程師（Q3 起），創辦人不支薪'],
        ['內容研發', '~NT$4 萬', '題庫、知識圖譜維護'],
        ['行銷推廣', '~NT$4 萬', '數位行銷、KOL、教育展'],
        ['營運行政', '~NT$2 萬', '辦公、法務、雜支'],
        ['**合計**', '**~NT$16 萬**', ''],
    ],
    [2200, 2200, 4626]
))

new_elements.append(make_body_text([
    ('關鍵特性', True, ACCENT_COLOR),
    ('：毛利率 80% 為固定值，不因規模變化。隨用戶數成長，固定成本被攤薄，', False, TEXT_COLOR),
    ('淨利率從 Year 1 的 -145% 快速改善至 Year 3 的 44%', True, TEXT_COLOR),
    ('——這是 SaaS 模式最核心的規模經濟效應。', False, TEXT_COLOR),
]))

# === Insert new elements before "3. 通路與品牌策略" ===
for el in new_elements:
    insert_before.addprevious(el)

doc.save(OUTPUT)
print(f'Saved to {OUTPUT}')
print(f'Inserted {len(new_elements)} new elements')
